"""Word (DOCX) rapor üretici modülü."""

import os
from pathlib import Path
from typing import Dict, Any, Optional
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    Document = None

import yaml
from rich.console import Console

from ..processor.structurer import StructuredReport, ReportSection

console = Console()


class DocxGenerator:
    """Word dokümanı üretici sınıfı."""

    # Türkçe rapor türü isimleri
    REPORT_TYPE_NAMES = {
        "is_plani": "İş Planı",
        "proje_raporu": "Proje Raporu",
        "sunum": "Sunum",
        "on_fizibilite": "Ön Fizibilite Raporu",
        "teknik_dok": "Teknik Dokümantasyon",
        "analiz_raporu": "Analiz Raporu",
        "kisa_not": "Kısa Not"
    }

    def __init__(self, config_dir: str = None):
        if Document is None:
            raise ImportError("python-docx kütüphanesi yüklü değil. 'pip install python-docx' komutunu çalıştırın.")

        self.config_dir = config_dir or Path(__file__).parent.parent.parent / "config"
        self.templates_dir = Path(__file__).parent.parent.parent / "templates"

        # Kuralları yükle
        self.rules = self._load_rules()
        self.styles = self._load_styles()

    def _load_rules(self) -> Dict[str, Any]:
        """Kuralları yükle."""
        rules_path = Path(self.config_dir) / "rules.yaml"
        if rules_path.exists():
            with open(rules_path, 'r', encoding='utf-8') as f:
                return yaml.safe_load(f)
        return {}

    def _load_styles(self) -> Dict[str, Any]:
        """Stilleri yükle."""
        styles_path = self.templates_dir / "styles.json"
        if styles_path.exists():
            import json
            with open(styles_path, 'r', encoding='utf-8') as f:
                return json.load(f)

        # Varsayılan stiller
        return {
            "colors": {
                "primary": "#1a365d",
                "secondary": "#2c5282",
                "accent": "#3182ce",
                "text": "#1a202c"
            },
            "fonts": {
                "heading": "Arial",
                "body": "Calibri"
            }
        }

    def _hex_to_rgb(self, hex_color: str) -> tuple:
        """Hex rengi RGB'ye çevir."""
        hex_color = hex_color.lstrip('#')
        return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))

    def _setup_styles(self, doc: Document):
        """Doküman stillerini ayarla."""
        fmt_rules = self.rules.get('format_rules', {})
        fonts = fmt_rules.get('fonts', {})
        colors = fmt_rules.get('colors', {})

        # Başlık stilleri
        for i in range(1, 4):
            style_name = f'Heading {i}'
            try:
                style = doc.styles[style_name]
            except KeyError:
                style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)

            font = style.font
            font.name = fonts.get('heading', {}).get('name', 'Arial')
            font.bold = True

            if i == 1:
                font.size = Pt(fonts.get('heading', {}).get('size_h1', 18))
                rgb = self._hex_to_rgb(colors.get('primary', '#1a365d'))
            elif i == 2:
                font.size = Pt(fonts.get('heading', {}).get('size_h2', 14))
                rgb = self._hex_to_rgb(colors.get('secondary', '#2c5282'))
            else:
                font.size = Pt(fonts.get('heading', {}).get('size_h3', 12))
                rgb = self._hex_to_rgb(colors.get('accent', '#3182ce'))

            font.color.rgb = RGBColor(*rgb)

        # Normal stil
        try:
            normal_style = doc.styles['Normal']
            normal_style.font.name = fonts.get('body', {}).get('name', 'Calibri')
            normal_style.font.size = Pt(fonts.get('body', {}).get('size', 11))
        except Exception:
            pass

    def _add_page_number(self, doc: Document):
        """Sayfa numarası ekle."""
        for section in doc.sections:
            footer = section.footer
            footer.is_linked_to_previous = False

            paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Sayfa numarası field'ı
            run = paragraph.add_run()
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            run._r.append(fldChar1)

            run2 = paragraph.add_run()
            instrText = OxmlElement('w:instrText')
            instrText.text = "PAGE"
            run2._r.append(instrText)

            run3 = paragraph.add_run()
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')
            run3._r.append(fldChar2)

    def _add_cover_page(self, doc: Document, report: StructuredReport):
        """Kapak sayfası ekle."""
        colors = self.rules.get('format_rules', {}).get('colors', {})
        primary_rgb = self._hex_to_rgb(colors.get('primary', '#1a365d'))

        # Boşluk
        for _ in range(5):
            doc.add_paragraph()

        # Başlık
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(report.title.upper())
        title_run.bold = True
        title_run.font.size = Pt(28)
        title_run.font.color.rgb = RGBColor(*primary_rgb)

        # Alt başlık (rapor türü)
        doc.add_paragraph()
        type_para = doc.add_paragraph()
        type_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        type_name = self.REPORT_TYPE_NAMES.get(report.report_type, report.report_type.replace('_', ' '))
        type_run = type_para.add_run(type_name)
        type_run.font.size = Pt(16)
        type_run.font.color.rgb = RGBColor(100, 100, 100)

        # Boşluk
        for _ in range(10):
            doc.add_paragraph()

        # Tarih
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_format = "DD.MM.YYYY" if report.language == "tr" else "YYYY-MM-DD"
        if report.language == "tr":
            date_str = datetime.now().strftime("%d.%m.%Y")
        else:
            date_str = datetime.now().strftime("%Y-%m-%d")
        date_run = date_para.add_run(date_str)
        date_run.font.size = Pt(12)

        # Sayfa sonu
        doc.add_page_break()

    def _add_table_of_contents(self, doc: Document, report: StructuredReport):
        """Gerçek Word İçindekiler ekle (sayfa numaralı)."""
        toc_title = "İÇİNDEKİLER" if report.language == "tr" else "TABLE OF CONTENTS"

        heading = doc.add_heading(toc_title, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        # Word TOC Field ekle - otomatik sayfa numarası
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()

        # TOC field başlangıcı
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar1)

        # TOC komutu - Heading 1-3 seviyelerini içerir
        run2 = paragraph.add_run()
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = ' TOC \\o "1-3" \\h \\z \\u '  # 1-3 seviye başlıklar, hyperlink, sayfa no
        run2._r.append(instrText)

        # TOC field sonu
        run3 = paragraph.add_run()
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        run3._r.append(fldChar2)

        # Placeholder metin (Word açıldığında güncellenecek)
        update_msg = "İçindekiler tablosunu güncellemek için sağ tıklayın ve 'Alanı Güncelle' seçin." if report.language == "tr" else "Right-click and select 'Update Field' to update table of contents."
        run4 = paragraph.add_run(update_msg)
        run4.italic = True
        run4.font.color.rgb = RGBColor(128, 128, 128)

        run5 = paragraph.add_run()
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        run5._r.append(fldChar3)

        doc.add_paragraph()
        doc.add_page_break()

    def _parse_markdown_table(self, text: str) -> list:
        """Markdown tablosunu parse et."""
        lines = [l.strip() for l in text.strip().split('\n') if l.strip()]
        if len(lines) < 2:
            return None

        # | ile başlayan satırları bul
        table_lines = [l for l in lines if l.startswith('|') and l.endswith('|')]
        if len(table_lines) < 2:
            return None

        rows = []
        for line in table_lines:
            # Ayırıcı satırı atla (| --- | --- |)
            if '---' in line:
                continue
            cells = [c.strip() for c in line.split('|')[1:-1]]
            if cells:
                rows.append(cells)

        return rows if len(rows) >= 1 else None

    def _add_table(self, doc: Document, rows: list):
        """Word tablosu ekle."""
        if not rows or len(rows) < 1:
            return

        table = doc.add_table(rows=len(rows), cols=len(rows[0]))
        table.style = 'Table Grid'

        for i, row_data in enumerate(rows):
            row = table.rows[i]
            for j, cell_text in enumerate(row_data):
                if j < len(row.cells):
                    cell = row.cells[j]
                    cell.text = str(cell_text)
                    # İlk satır başlık - kalın yap
                    if i == 0:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True

        doc.add_paragraph()  # Tablo sonrası boşluk

    def _add_section(self, doc: Document, section: ReportSection, number: str):
        """Bölüm ekle."""
        # Başlık
        heading_text = f"{number} {section.title}"
        doc.add_heading(heading_text, level=section.level)

        # İçerik paragrafları
        if section.content:
            paragraphs = section.content.split('\n\n')
            for para_text in paragraphs:
                para_text = para_text.strip()
                if not para_text:
                    continue

                # Markdown tablo kontrolü
                if '|' in para_text and para_text.count('|') >= 4:
                    table_rows = self._parse_markdown_table(para_text)
                    if table_rows:
                        self._add_table(doc, table_rows)
                        continue

                # Alt başlık kontrolü
                if para_text.startswith('## '):
                    doc.add_heading(para_text[3:], level=2)
                elif para_text.startswith('### '):
                    doc.add_heading(para_text[4:], level=3)
                elif para_text.startswith('[Kaynak:'):
                    # Kaynak notunu italik yap
                    p = doc.add_paragraph()
                    run = p.add_run(para_text)
                    run.italic = True
                    run.font.size = Pt(9)
                elif para_text.startswith('- ') or para_text.startswith('* '):
                    # Madde işaretli liste
                    for line in para_text.split('\n'):
                        line = line.strip()
                        if line.startswith('- ') or line.startswith('* '):
                            doc.add_paragraph(line[2:], style='List Bullet')
                        elif line.startswith('  - ') or line.startswith('  * '):
                            # Alt madde
                            doc.add_paragraph(line[4:], style='List Bullet')
                        elif line:
                            doc.add_paragraph(line)
                elif len(para_text) > 2 and para_text[0].isdigit() and (para_text[1] == '.' or para_text[1] == ')'):
                    # Numaralı liste
                    for line in para_text.split('\n'):
                        line = line.strip()
                        if line and len(line) > 2 and line[0].isdigit():
                            text = line.lstrip('0123456789.)').strip()
                            if text.startswith(':'):
                                text = text[1:].strip()
                            doc.add_paragraph(text, style='List Number')
                        elif line.strip().startswith('- '):
                            # Alt madde
                            doc.add_paragraph(line.strip()[2:], style='List Bullet')
                        elif line:
                            doc.add_paragraph(line)
                else:
                    # Normal paragraf
                    doc.add_paragraph(para_text)

        # Alt bölümler
        for i, subsection in enumerate(section.subsections, 1):
            sub_number = f"{number}.{i}"
            self._add_section(doc, subsection, sub_number)

    def generate(self, report: StructuredReport, output_path: str) -> str:
        """Raporu Word dosyası olarak oluştur."""
        doc = Document()

        # Stilleri ayarla
        self._setup_styles(doc)

        # Sayfa kenar boşlukları
        for section in doc.sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

        # Kapak sayfası
        self._add_cover_page(doc, report)

        # İçindekiler
        self._add_table_of_contents(doc, report)

        # Bölümler
        section_num = 1
        for section in report.sections:
            # Kapak ve içindekiler bölümlerini atla
            if section.id in ['kapak', 'kapak_sayfasi', 'icindekiler', 'table_of_contents']:
                continue

            self._add_section(doc, section, str(section_num))
            section_num += 1

        # Sayfa numaraları
        self._add_page_number(doc)

        # Dosya adı oluştur
        output_path = Path(output_path)
        if output_path.is_dir():
            filename = f"{report.report_type}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            output_path = output_path / filename

        # Kaydet
        doc.save(str(output_path))

        return str(output_path)
