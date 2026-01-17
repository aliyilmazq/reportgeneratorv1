"""Word dosya okuyucu modülü."""

from pathlib import Path
from typing import Dict, Any, List, Optional
from dataclasses import dataclass, field

try:
    from docx import Document
    from docx.table import Table as DocxTable
    from docx.text.paragraph import Paragraph
except ImportError:
    Document = None


@dataclass
class WordTable:
    """Word'den çıkarılan tablo."""
    data: List[List[str]]
    headers: List[str] = field(default_factory=list)
    row_count: int = 0
    column_count: int = 0


@dataclass
class WordParagraph:
    """Word paragrafı."""
    text: str
    style: str  # 'Heading 1', 'Normal', etc.
    is_heading: bool = False
    heading_level: int = 0


@dataclass
class WordContent:
    """Word dosyası içeriği."""
    source: str
    filename: str
    paragraphs: List[WordParagraph] = field(default_factory=list)
    tables: List[WordTable] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)
    full_text: str = ""
    image_count: int = 0


class WordParser:
    """Word dosya okuyucu sınıfı."""

    def __init__(self):
        if Document is None:
            raise ImportError("python-docx kütüphanesi yüklü değil. 'pip install python-docx' komutunu çalıştırın.")

    def parse(self, file_path: str) -> WordContent:
        """Word dosyasını oku ve içeriği çıkar."""
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"Dosya bulunamadı: {file_path}")

        content = WordContent(
            source=str(file_path),
            filename=file_path.name
        )

        try:
            doc = Document(str(file_path))

            # Metadata
            core_props = doc.core_properties
            content.metadata = {
                'title': core_props.title or "",
                'author': core_props.author or "",
                'created': str(core_props.created) if core_props.created else "",
                'modified': str(core_props.modified) if core_props.modified else "",
                'subject': core_props.subject or ""
            }

            # Paragrafları oku
            all_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    style_name = para.style.name if para.style else "Normal"
                    is_heading = style_name.startswith("Heading")
                    heading_level = 0

                    if is_heading:
                        try:
                            heading_level = int(style_name.split()[-1])
                        except (ValueError, IndexError):
                            heading_level = 1

                    word_para = WordParagraph(
                        text=para.text.strip(),
                        style=style_name,
                        is_heading=is_heading,
                        heading_level=heading_level
                    )
                    content.paragraphs.append(word_para)
                    all_text.append(para.text.strip())

            content.full_text = "\n\n".join(all_text)

            # Tabloları oku
            for table in doc.tables:
                word_table = self._parse_table(table)
                if word_table:
                    content.tables.append(word_table)

            # Görsel sayısını bul
            content.image_count = self._count_images(doc)

        except Exception as e:
            content.metadata['error'] = str(e)

        return content

    def _parse_table(self, table: 'DocxTable') -> Optional[WordTable]:
        """Word tablosunu parse et."""
        try:
            rows = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                rows.append(row_data)

            if not rows:
                return None

            # İlk satırı başlık olarak al
            headers = rows[0] if rows else []
            data = rows[1:] if len(rows) > 1 else []

            return WordTable(
                headers=headers,
                data=data,
                row_count=len(rows),
                column_count=len(headers)
            )
        except Exception:
            return None

    def _count_images(self, doc) -> int:
        """Dokümandaki görsel sayısını bul."""
        count = 0
        try:
            for rel in doc.part.rels.values():
                if "image" in rel.reltype:
                    count += 1
        except Exception:
            pass
        return count

    def get_headings(self, content: WordContent) -> List[Dict[str, Any]]:
        """Başlıkları hiyerarşik olarak al."""
        headings = []
        for para in content.paragraphs:
            if para.is_heading:
                headings.append({
                    'level': para.heading_level,
                    'text': para.text
                })
        return headings

    def get_text_by_heading(self, content: WordContent) -> Dict[str, str]:
        """İçeriği başlıklara göre grupla."""
        sections = {}
        current_heading = "Başlıksız"
        current_text = []

        for para in content.paragraphs:
            if para.is_heading:
                # Önceki bölümü kaydet
                if current_text:
                    sections[current_heading] = "\n".join(current_text)
                current_heading = para.text
                current_text = []
            else:
                current_text.append(para.text)

        # Son bölümü kaydet
        if current_text:
            sections[current_heading] = "\n".join(current_text)

        return sections

    def to_dict(self, content: WordContent) -> Dict[str, Any]:
        """WordContent'i sözlüğe çevir."""
        return {
            "type": "word",
            "source": content.source,
            "filename": content.filename,
            "metadata": content.metadata,
            "full_text": content.full_text,
            "paragraphs": [
                {
                    "text": para.text,
                    "style": para.style,
                    "is_heading": para.is_heading,
                    "heading_level": para.heading_level
                }
                for para in content.paragraphs
            ],
            "tables": [
                {
                    "headers": table.headers,
                    "data": table.data,
                    "row_count": table.row_count,
                    "column_count": table.column_count
                }
                for table in content.tables
            ],
            "image_count": content.image_count
        }

    def to_markdown(self, content: WordContent) -> str:
        """Word içeriğini Markdown formatına çevir."""
        markdown = []

        for para in content.paragraphs:
            if para.is_heading:
                prefix = "#" * para.heading_level
                markdown.append(f"{prefix} {para.text}\n")
            else:
                markdown.append(f"{para.text}\n")

        # Tabloları ekle
        for i, table in enumerate(content.tables, 1):
            markdown.append(f"\n**Tablo {i}:**\n")
            if table.headers:
                header_row = "| " + " | ".join(table.headers) + " |"
                separator = "| " + " | ".join(["---"] * len(table.headers)) + " |"
                markdown.append(header_row)
                markdown.append(separator)

                for row in table.data:
                    row_str = "| " + " | ".join(str(cell)[:50] for cell in row) + " |"
                    markdown.append(row_str)
            markdown.append("")

        return "\n".join(markdown)
