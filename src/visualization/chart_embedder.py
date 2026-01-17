"""Grafik Gömme Modülü - Grafikleri Word/PDF'e gömer."""

import os
from typing import List, Dict, Any, Optional
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    Document = None

from rich.console import Console

from .chart_generator import GeneratedChart

console = Console()


class ChartEmbedder:
    """Grafik gömücü - Word ve PDF'e grafik ekler."""

    def __init__(self):
        pass

    def embed_in_docx(
        self,
        doc: 'Document',
        chart: GeneratedChart,
        width_inches: float = 5.5,
        caption: str = None,
        position: str = "center"
    ) -> bool:
        """Word dökümanına grafik göm."""
        if Document is None:
            console.print("[yellow]python-docx yüklü değil[/yellow]")
            return False

        if not os.path.exists(chart.file_path):
            console.print(f"[yellow]Grafik dosyası bulunamadı: {chart.file_path}[/yellow]")
            return False

        try:
            # Grafik için paragraf
            paragraph = doc.add_paragraph()

            # Hizalama
            if position == "center":
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif position == "right":
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # Grafik ekle
            run = paragraph.add_run()
            run.add_picture(chart.file_path, width=Inches(width_inches))

            # Başlık/açıklama ekle
            if caption or chart.title:
                caption_para = doc.add_paragraph()
                caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                caption_text = caption or f"Şekil: {chart.title}"
                caption_run = caption_para.add_run(caption_text)
                caption_run.italic = True
                caption_run.font.size = Pt(10)

            # Boşluk
            doc.add_paragraph()

            return True

        except Exception as e:
            console.print(f"[red]Grafik ekleme hatası: {e}[/red]")
            return False

    def embed_multiple_in_docx(
        self,
        doc: 'Document',
        charts: List[GeneratedChart],
        charts_per_row: int = 1,
        width_inches: float = 5.5
    ) -> int:
        """Birden fazla grafiği Word'e göm."""
        embedded_count = 0

        for chart in charts:
            if self.embed_in_docx(doc, chart, width_inches):
                embedded_count += 1

        return embedded_count

    def embed_charts_by_section(
        self,
        doc: 'Document',
        charts: List[GeneratedChart],
        section_chart_map: Dict[str, List[str]]
    ) -> int:
        """Grafikleri ilgili bölümlere göm."""
        embedded_count = 0

        for chart in charts:
            # Grafik başlığından bölüm belirle
            for section_id, chart_titles in section_chart_map.items():
                if any(t.lower() in chart.title.lower() for t in chart_titles):
                    if self.embed_in_docx(doc, chart):
                        embedded_count += 1
                    break

        return embedded_count

    def create_figure_list(self, charts: List[GeneratedChart]) -> str:
        """Şekil listesi oluştur."""
        if not charts:
            return ""

        lines = ["## Şekil Listesi\n"]

        for i, chart in enumerate(charts, 1):
            lines.append(f"Şekil {i}: {chart.title}")

        return "\n".join(lines)

    def get_chart_for_section(
        self,
        charts: List[GeneratedChart],
        section_id: str
    ) -> List[GeneratedChart]:
        """Bölüme ait grafikleri getir."""
        section_charts = []

        for chart in charts:
            # Dosya adı veya başlıkta bölüm ID'si var mı?
            if section_id.lower() in chart.file_path.lower() or \
               section_id.lower() in chart.title.lower():
                section_charts.append(chart)

        return section_charts


class PDFChartEmbedder:
    """PDF için grafik gömücü (ReportLab)."""

    def __init__(self):
        pass

    def embed_in_pdf(
        self,
        canvas,
        chart: GeneratedChart,
        x: float,
        y: float,
        width: float = 400,
        height: float = 300
    ) -> bool:
        """PDF'e grafik göm."""
        try:
            from reportlab.lib.utils import ImageReader

            if not os.path.exists(chart.file_path):
                return False

            img = ImageReader(chart.file_path)
            canvas.drawImage(img, x, y, width=width, height=height)
            return True

        except Exception as e:
            console.print(f"[red]PDF grafik hatası: {e}[/red]")
            return False
